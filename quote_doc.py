"""Fill the moving-quote Word form (`quote_template.docx`) and save a finished
견적서 `.docx`.

The form is a real estate / moving quote sheet (the 양식 the user supplied). We
keep it pristine in the repo and fill it *structurally* (by label), so it stays
editable in Word and survives small layout edits. Pipeline-derived fields
(운반차량 = chosen truck, 합계 = tonnage quote, 보관짐 = detected items + boxes) are
filled automatically; the customer / schedule fields come from the caller (the
web form) and are left blank when not given.

Only dependency is `python-docx`. `build_quote()` never invents data — any field
absent from `data` stays empty.
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

TEMPLATE = Path(__file__).with_name("quote_template.docx")


# --------------------------------------------------------------------------- #
# low-level helpers (work on any nested-table python-docx structure)
# --------------------------------------------------------------------------- #
def _iter_paragraphs(doc):
    """Every paragraph in document order, recursing into (nested) table cells."""
    def walk_table(tbl):
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nt in cell.tables:
                    yield from walk_table(nt)
    for tbl in doc.tables:
        yield from walk_table(tbl)
    for p in doc.paragraphs:
        yield p


def _iter_rows(doc):
    """Every (table, row) recursing into nested tables — for label→value cells."""
    def walk_table(tbl):
        for row in tbl.rows:
            yield tbl, row
            for cell in row.cells:
                for nt in cell.tables:
                    yield from walk_table(nt)
    for tbl in doc.tables:
        yield from walk_table(tbl)


def _append_run(paragraph, text):
    """Append `text` to a paragraph, inheriting the last run's formatting."""
    if not text:
        return
    runs = paragraph.runs
    if runs:                                   # clone last run's rPr for styling
        new = copy.deepcopy(runs[-1]._r)
        for child in list(new):                # drop its text node(s)
            if child.tag.endswith("}t"):
                new.remove(child)
        runs[-1]._r.addnext(new)
        r = Run(new, paragraph)
    else:
        r = paragraph.add_run("")
    r.text = text


def _set_para_text(paragraph, text):
    """Set a paragraph's text into its first run (clearing the rest)."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        paragraph.add_run(text)


def _clone_line(model_para, text):
    """A new <w:p> cloned from `model_para` (keeps font/size) carrying `text`."""
    new_p = copy.deepcopy(model_para._p)
    para = Paragraph(new_p, model_para._parent)
    _set_para_text(para, text)
    return para


def _set_cell(cell, text):
    """Replace a cell's text with `text` (keeps first paragraph's run style)."""
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for extra in p.runs[1:]:
            extra.text = ""
    else:
        p.add_run(text)


def _fill_prefix(paras, prefix, value, occurrence=0):
    """Append `value` to the `occurrence`-th paragraph whose text starts with prefix."""
    if value is None or value == "":
        return
    seen = 0
    for p in paras:
        if p.text.strip().startswith(prefix):
            if seen == occurrence:
                _append_run(p, str(value))
                return
            seen += 1


def _fill_label_cell(doc, label, value):
    """Find a cell whose text == label and fill the next non-label cell in its row."""
    if value is None or value == "":
        return
    for _tbl, row in _iter_rows(doc):
        cells = row.cells
        for i, c in enumerate(cells):
            if c.text.strip() == label:
                for j in range(i + 1, len(cells)):
                    if cells[j].text.strip() in ("", label):
                        if cells[j].text.strip() == "":
                            _set_cell(cells[j], str(value))
                            return
                return


def _fill_cost_cell(doc, label, value):
    """비용요약: the value cell already reads '만원'; put the number before it."""
    if value is None or value == "":
        return
    for _tbl, row in _iter_rows(doc):
        cells = row.cells
        for i, c in enumerate(cells):
            if c.text.strip() == label and i + 1 < len(cells):
                vc = cells[i + 1]
                if "만원" in vc.text:
                    p = vc.paragraphs[0]
                    if p.runs:
                        p.runs[0].text = f"{value}만원"
                        for extra in p.runs[1:]:
                            extra.text = ""
                    else:
                        p.add_run(f"{value}만원")
                else:
                    _set_cell(vc, f"{value}만원")
                return


def _body_cell(doc, anchor):
    """The body cell of the section whose header row contains `anchor`."""
    rows = list(_iter_rows(doc))
    for idx, (_tbl, row) in enumerate(rows):
        if any(anchor in c.text for c in row.cells) and idx + 1 < len(rows):
            return rows[idx + 1][1].cells[0]
    return None


def _append_lines_to_cell(doc, anchor, lines):
    """Append item lines to a section body cell, reusing its empty paragraphs first.

    Used for 세부비용 / 고객요청사항 — sections whose body cell is the row right after
    the header row. Empty template paragraphs are filled before new ones are added,
    so there are no stray blank lines between items.
    """
    if not lines:
        return
    body = _body_cell(doc, anchor)
    if body is None:
        return
    empties = [p for p in body.paragraphs if not p.text.strip()]
    model = body.paragraphs[0]
    for i, ln in enumerate(lines):
        if i < len(empties):
            _set_para_text(empties[i], str(ln))
        else:
            body._element.append(_clone_line(model, str(ln))._p)


def _fill_storage(doc, items, disposal):
    """보관 짐 / 폐기 짐: detected inventory goes ABOVE the '폐기짐' label, disposal below.

    The template body lists room labels ([안방]…) then '폐기짐(폐기장소이동)' then blank
    lines. We insert our detected items just before the 폐기짐 label (so they read as
    items to move, not to discard) and put disposal items after it.
    """
    body = _body_cell(doc, "보 관 짐 / 폐 기 짐")
    if body is None:
        return
    disposal_para = None
    for p in body.paragraphs:
        if p.text.strip().startswith("폐기짐"):
            disposal_para = p
            break
    model = body.paragraphs[0]
    for ln in (items or []):
        line = _clone_line(model, str(ln))
        if disposal_para is not None:
            disposal_para._p.addprevious(line._p)
        else:
            body._element.append(line._p)
    if disposal:
        empties = [p for p in body.paragraphs if not p.text.strip()]
        for i, ln in enumerate(disposal):
            if i < len(empties):
                _set_para_text(empties[i], str(ln))
            else:
                body._element.append(_clone_line(model, str(ln))._p)


# --------------------------------------------------------------------------- #
# public API
# --------------------------------------------------------------------------- #
def build_quote(data: dict, out_path: str, template: str | None = None) -> str:
    """Fill the quote template with `data` and write a .docx to `out_path`.

    Recognised keys (all optional; missing → left blank):
      issue_date, valid_until, manager,
      customer_name, customer_phone,
      origin_address, origin_date, origin_time,
      dest_address, dest_date, dest_time,
      work_in_type, work_out_type, crew_in, crew_out, vehicle,
      cost_in, cost_out, deposit, total,            # 만원 (numbers/str)
      storage_period, storage_total, storage_daily,
      detail_cost (list[str]), requests (list[str]),
      items (list[str]), disposal (list[str]).
    Returns `out_path`.
    """
    doc = Document(template or str(TEMPLATE))
    paras = list(_iter_paragraphs(doc))

    # header
    _fill_prefix(paras, "발행일:", data.get("issue_date"))
    _fill_prefix(paras, "유효기간:", data.get("valid_until"))
    _fill_prefix(paras, "담당자:", data.get("manager"))

    # 고객 정보
    _fill_label_cell(doc, "고객명", data.get("customer_name"))
    _fill_label_cell(doc, "연락처", data.get("customer_phone"))

    # 출발지 / 도착지 (주소 appears twice: origin then destination)
    _fill_prefix(paras, "주소", data.get("origin_address"), occurrence=0)
    _fill_prefix(paras, "주소", data.get("dest_address"), occurrence=1)
    _fill_prefix(paras, "입고일자", data.get("origin_date"))
    _fill_prefix(paras, "입고 시작", data.get("origin_time"))
    _fill_prefix(paras, "출고일자", data.get("dest_date"))
    _fill_prefix(paras, "출고 시작", data.get("dest_time"))

    # 작업 정보
    _fill_label_cell(doc, "입고 작업유형", data.get("work_in_type"))
    _fill_label_cell(doc, "출고 작업유형", data.get("work_out_type"))
    _fill_label_cell(doc, "입고 인원(남:녀)", data.get("crew_in"))
    _fill_label_cell(doc, "출고 인원(남:녀)", data.get("crew_out"))
    _fill_label_cell(doc, "운반차량", data.get("vehicle"))

    # 비용 요약 (단위: 만원)
    _fill_cost_cell(doc, "입고 총액", data.get("cost_in"))
    _fill_cost_cell(doc, "출고 총액", data.get("cost_out"))
    _fill_cost_cell(doc, "계약금", data.get("deposit"))
    _fill_cost_cell(doc, "합계", data.get("total"))

    # 보관비 요약 (line: '보관기간  start - end  · 보관비 총액  N 만원')
    if data.get("storage_period") or data.get("storage_total") or data.get("storage_daily"):
        for p in paras:
            if p.text.strip().startswith("보관기간"):
                runs = p.runs
                # runs ≈ ['보관기간 ', '   -   ', ' · 보관비 총액 ', '   만원']
                if data.get("storage_period") and len(runs) >= 2:
                    runs[1].text = f" {data['storage_period']} "
                if data.get("storage_total") and len(runs) >= 4:
                    runs[3].text = f" {data['storage_total']}만원"
                break
        if data.get("storage_daily"):
            for p in paras:
                if "일자별 정산" in p.text and p.runs:
                    p.runs[0].text = p.runs[0].text.replace(
                        "1일         원", f"1일 {data['storage_daily']}원")
                    break

    # free-text sections
    _append_lines_to_cell(doc, "세 부 비 용", data.get("detail_cost"))
    _append_lines_to_cell(doc, "고 객 요 청 사 항", data.get("requests"))

    # 보관 짐 / 폐기 짐 — detected inventory above the 폐기짐 label, disposal below
    _fill_storage(doc, data.get("items"), data.get("disposal"))

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


# --------------------------------------------------------------------------- #
# build the data dict from a pipeline result (viz JSON)
# --------------------------------------------------------------------------- #
def data_from_viz(viz: dict, form: dict | None = None) -> dict:
    """Merge a pipeline viz dict with the VLM quote estimate and the web-form fields.

    Precedence (low → high):
      1. quote_estimate — the VLM's predicted 작업유형 / 인원 / 비용 / 요청사항 / 합계
         (so the quote is fully pre-filled after a run; the human just edits).
      2. pipeline-auto — vehicle (truck), detail_cost (truck breakdown), items
         (kept objects + boxes); these are authoritative from the run.
      3. form — the human's edits in the UI win over everything.
    """
    form = dict(form or {})
    data = dict(viz.get("quote_estimate") or {})     # 1. VLM estimate as the base
    data.pop("source", None)

    tp = viz.get("truck_plan") or {}
    truck = tp.get("truck") or {}
    if truck.get("name"):                            # 2. pipeline-authoritative
        data["vehicle"] = _vehicle_label(truck, tp.get("count", 1))
    quote = tp.get("quote_krw")
    if quote and "total" not in data:                # estimate already set total; else truck fare
        data["total"] = _man(quote)

    # 세부비용: truck line(s)
    detail = list(form.get("detail_cost") or [])
    if truck.get("name") and not detail:
        cnt = tp.get("count", 1)
        fare = truck.get("price_krw", 0)
        detail.append(f"{truck['name']} x {cnt}대 = {_man(cnt * fare)}만원 (차량+인력)")
        if tp.get("cargo_volume_m3"):
            detail.append(f"총 화물 부피 {tp['cargo_volume_m3']} m3, "
                          f"적재율 {round(tp.get('utilization', 0) * 100)}%")
    if detail:
        data["detail_cost"] = detail
    data["items"] = _inventory_lines(viz)

    data.update(form)                                # 3. human edits win
    return data


def _man(krw) -> str:
    """KRW → 만원 string (10000원 = 1만원), trimmed."""
    v = krw / 10000.0
    return str(int(round(v))) if abs(v - round(v)) < 1e-6 else f"{v:.1f}"


def _vehicle_label(truck: dict, count: int) -> str:
    name = truck.get("name", "")
    return f"{name} x {count}대" if count and count > 1 else name


def _inventory_lines(viz: dict) -> list[str]:
    """One line per kept object: identity + measured size + assigned box."""
    lines = []
    box_count: dict[str, int] = {}
    for o in viz.get("instances", []):
        if not o.get("kept", True):
            continue
        name = o.get("identity") or o.get("label") or f"object {o.get('id')}"
        dims = o.get("corrected_dims_cm") or o.get("dims_cm") or []
        size = "x".join(str(round(float(d))) for d in dims) + "cm" if dims else ""
        box = (o.get("box") or "").strip()
        where = f"-> {box}" if box else "(loose 가구)"
        lines.append(f"{name} {size} {where}".strip())
        if box:
            box_count[box] = box_count.get(box, 0) + 1
    if box_count:
        summary = ", ".join(f"{b} {n}개" for b, n in sorted(box_count.items()))
        lines.append(f"[박스 합계] {summary}")
    return lines
